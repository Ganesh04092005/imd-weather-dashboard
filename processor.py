import pandas as pd
from collections import defaultdict
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# 🔹 FOOTER (REPEATS EVERY PAGE)
def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer

    p = footer.paragraphs[0]

    p.text = (
        "Contact: India Meteorological Department, Meteorological Centre Hyderabad.\n"
        "Phone: (91) 040-27908506, FAX: (91) 040-27906172\n"
        "Website: imdhyderabad.imd.gov.in\n"
        "Spatial rainfall distribution: Isolated <25%, Few 26-50%, Many 51-75%, Most 76-100%\n"
        "Rainfall: Light 2.5–15.5mm, Moderate 15.6–64.4mm, Heavy 64.5–115.5mm,\n"
        "Very Heavy 115.6–204.4mm, Extremely Heavy >204.5mm\n"
        "Follow us: Mausam App | Meghdoot App | Damini App"
    )

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# 🔹 TABLE BORDER FIX (IMPORTANT)
def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)


# 🔹 Convert IMD format
def convert_imd(file):
    df = pd.read_excel(file, header=None)

    data = []

    for i in range(len(df)):
        try:
            district = df.iloc[i, 2]

            if pd.isna(district):
                continue

            row = {
                "DISTRICT": str(district).strip(),

                "DAY1_wrng": df.iloc[i, 5] if df.shape[1] > 5 else "",
                "DAY2_wrng": df.iloc[i, 8] if df.shape[1] > 8 else "",
                "DAY3_wrng": df.iloc[i, 11] if df.shape[1] > 11 else "",
                "DAY4_wrng": df.iloc[i, 14] if df.shape[1] > 14 else "",
                "DAY5_wrng": df.iloc[i, 17] if df.shape[1] > 17 else "",
                "DAY6_wrng": df.iloc[i, 19] if df.shape[1] > 19 else "",
                "DAY7_wrng": df.iloc[i, 20] if df.shape[1] > 20 else "",
            }

            data.append(row)

        except:
            continue

    return pd.DataFrame(data)


# 🔹 Classification
def classify(text):
    text = str(text)

    if "Extremely Heavy" in text:
        return "EXTREME"
    elif "Very Heavy" in text:
        return "VERY_HEAVY"
    elif "Heavy" in text:
        return "HEAVY"
    else:
        return "NORMAL"


# 🔹 Process
def process(df, day):
    grouped = defaultdict(list)

    for _, row in df.iterrows():
        district = row["DISTRICT"]
        warning = classify(row[f"DAY{day}_wrng"])
        grouped[warning].append(district)

    return grouped


# 🔹 Cell color
def set_cell_color(cell, color):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


# 🔹 FINAL DOCUMENT GENERATION
from docxtpl import DocxTemplate
from datetime import datetime, timedelta


# from docxtpl import DocxTemplate
# from datetime import datetime, timedelta


# 🔥 Clean bullet formatter
def format_bullets(items):
    return "\n".join([f"• {i}" for i in items])


def generate_doc(file):
    df = convert_imd(file)

    doc = DocxTemplate(r"C:\Users\ganes\OneDrive\Desktop\IMD PROJECT\imd_template.docx")

    context = {}

    today = datetime.today()

    # 🔷 ISSUE DATE
    context["ISSUE_DATE"] = today.strftime("%d-%m-%Y")

    for day in range(1, 8):

        # 🔷 DATE CALCULATION
        from_date = today + timedelta(days=day - 1)
        to_date = today + timedelta(days=day)

        context[f"DAY{day}_FROM"] = from_date.strftime("%d/%m/%Y")
        context[f"DAY{day}_TO"] = to_date.strftime("%d/%m/%Y")

        grouped = process(df, day)

        # 🔷 FORECAST
        forecast = (
            "Light to Moderate Rain or Thundershowers very likely "
            "to occur at isolated/few places over Telangana."
        )

        # 🔷 WARNING (FIXED FORMAT)
        warning = ""

        if grouped["EXTREME"]:
            warning += "Very Heavy to Extremely Heavy Rainfall:\n"
            warning += format_bullets(grouped["EXTREME"]) + "\n\n"

        if grouped["VERY_HEAVY"]:
            warning += "Heavy to Very Heavy Rainfall:\n"
            warning += format_bullets(grouped["VERY_HEAVY"]) + "\n\n"

        if grouped["HEAVY"]:
            warning += "Heavy Rainfall:\n"
            warning += format_bullets(grouped["HEAVY"])

        if warning.strip() == "":
            warning = "NIL"

        # 🔷 PASS TO TEMPLATE
        context[f"DAY{day}_FORECAST"] = forecast
        context[f"DAY{day}_WARNING"] = warning

    # 🔷 RENDER TEMPLATE
    doc.render(context)

    output_file = "FINAL_IMD_OUTPUT.docx"
    doc.save(output_file)

    return output_file


# def generate_doc(file):
#     df = convert_imd(file)

#     doc = Document()

#     # Add footer
#     add_footer(doc)

#     # HEADER
#     header = doc.add_paragraph()
#     run = header.add_run(
#         "Government of India\n"
#         "India Meteorological Department\n"
#         "(Ministry of Earth Sciences)\n"
#         "Meteorological Centre, Hyderabad\n"
#     )
#     run.bold = True
#     run.font.size = Pt(12)
#     header.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     doc.add_paragraph()

#     # TITLE
#     title = doc.add_paragraph()
#     t = title.add_run(
#         "Seven Day Forecast and Farmer’s Weather Bulletin for TELANGANA STATE\n"
#         "ISSUED AT 1300 HRS IST (MID-DAY)"
#     )
#     t.bold = True
#     title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     doc.add_paragraph("-" * 80)

#     # MAIN TABLE
#     table = doc.add_table(rows=1, cols=3)
#     table.alignment = WD_TABLE_ALIGNMENT.CENTER
#     table.style = "Table Grid"
#     set_table_borders(table)

#     header_cells = table.rows[0].cells
#     header_cells[0].text = "DAY"
#     header_cells[1].text = "FORECAST"
#     header_cells[2].text = "WARNING"

#     for day in range(1, 8):
#         grouped = process(df, day)

#         row = table.add_row().cells

#         row[0].text = f"DAY {day}"
#         row[1].text = "Light to Moderate Rain or Thundershowers likely."

#         warning_text = ""

#         if grouped["EXTREME"]:
#             warning_text += "Very Heavy to Extremely Heavy Rainfall:\n" + ", ".join(grouped["EXTREME"]) + "\n\n"

#         if grouped["VERY_HEAVY"]:
#             warning_text += "Heavy to Very Heavy Rainfall:\n" + ", ".join(grouped["VERY_HEAVY"]) + "\n\n"

#         if grouped["HEAVY"]:
#             warning_text += "Heavy Rainfall:\n" + ", ".join(grouped["HEAVY"]) + "\n"

#         row[2].text = warning_text

#         if grouped["EXTREME"]:
#             set_cell_color(row[2], "FF0000")
#         elif grouped["VERY_HEAVY"]:
#             set_cell_color(row[2], "FFA500")
#         elif grouped["HEAVY"]:
#             set_cell_color(row[2], "FFFF00")

#     # SECOND PAGE
#     doc.add_page_break()

#     impact_table = doc.add_table(rows=4, cols=3)
#     impact_table.style = "Table Grid"
#     set_table_borders(impact_table)

#     impact_table.rows[0].cells[0].text = "RISK"
#     impact_table.rows[0].cells[1].text = "IMPACT"
#     impact_table.rows[0].cells[2].text = "ACTION"

#     impact_table.rows[1].cells[0].text = "MEDIUM"
#     impact_table.rows[2].cells[0].text = "HIGH"
#     impact_table.rows[3].cells[0].text = "VERY HIGH"

#     set_cell_color(impact_table.rows[1].cells[0], "FFFF00")
#     set_cell_color(impact_table.rows[2].cells[0], "FFA500")
#     set_cell_color(impact_table.rows[3].cells[0], "FF0000")

#     # FOOTER SIGNATURE
#     doc.add_paragraph("\nDirector I/C        Duty Officer")

#     output_file = "FINAL_IMD_BULLETIN.docx"
#     doc.save(output_file)

#     return output_file
# from datetime import datetime, timedelta


# def generate_doc(file):
#     df = convert_imd(file)

#     doc = Document()

#     # Add footer
#     add_footer(doc)

#     # 🔷 HEADER
#     header = doc.add_paragraph()
#     run = header.add_run(
#         "Government of India\n"
#         "India Meteorological Department\n"
#         "(Ministry of Earth Sciences)\n"
#         "Meteorological Centre, Hyderabad\n"
#     )
#     run.bold = True
#     run.font.size = Pt(12)
#     header.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     doc.add_paragraph()

#     # 🔷 DATE (DYNAMIC)
#     today = datetime.today()
#     issue_date = today.strftime("%d-%m-%Y")

#     # 🔷 TITLE WITH DATE
#     title = doc.add_paragraph()
#     t = title.add_run(
#         f"Seven Day Forecast and Farmer’s Weather Bulletin for TELANGANA STATE\n"
#         f"ISSUED AT 1300 HRS IST (MID-DAY) ON {issue_date}"
#     )
#     t.bold = True
#     title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     doc.add_paragraph("-" * 80)

#     # 🔷 MAIN TABLE
#     table = doc.add_table(rows=1, cols=4)
#     table.alignment = WD_TABLE_ALIGNMENT.CENTER
#     table.style = "Table Grid"
#     set_table_borders(table)

#     header_cells = table.rows[0].cells
#     header_cells[0].text = "DAY"
#     header_cells[1].text = "PERIOD"
#     header_cells[2].text = "FORECAST"
#     header_cells[3].text = "WARNING"

#     for day in range(1, 8):
#         grouped = process(df, day)

#         # 🔷 DATES
#         from_date = today + timedelta(days=day - 1)
#         to_date = today + timedelta(days=day)

#         from_str = from_date.strftime("%d/%m/%Y")
#         to_str = to_date.strftime("%d/%m/%Y")

#         row = table.add_row().cells

#         row[0].text = f"DAY {day}"
#         row[1].text = f"From: {from_str}\nTo: {to_str}"

#         # 🔷 FORECAST (IMD STYLE)
#         forecast = (
#             "Light to Moderate Rain or Thundershowers very likely "
#             "to occur at isolated/few places over Telangana."
#         )

#         row[2].text = forecast

#         # 🔷 WARNING (BULLET FORMAT)
#         warning_text = ""

#         if grouped["EXTREME"]:
#             warning_text += "Very Heavy to Extremely Heavy Rainfall:\n"
#             warning_text += "\n• " + "\n• ".join(grouped["EXTREME"]) + "\n\n"

#         if grouped["VERY_HEAVY"]:
#             warning_text += "Heavy to Very Heavy Rainfall:\n"
#             warning_text += "\n• " + "\n• ".join(grouped["VERY_HEAVY"]) + "\n\n"

#         if grouped["HEAVY"]:
#             warning_text += "Heavy Rainfall:\n"
#             warning_text += "\n• " + "\n• ".join(grouped["HEAVY"])

#         if warning_text.strip() == "":
#             warning_text = "NIL"

#         row[3].text = warning_text

#         # 🔷 COLORS
#         if grouped["EXTREME"]:
#             set_cell_color(row[3], "FF0000")
#         elif grouped["VERY_HEAVY"]:
#             set_cell_color(row[3], "FFA500")
#         elif grouped["HEAVY"]:
#             set_cell_color(row[3], "FFFF00")

#     # 🔷 SECOND PAGE
#     doc.add_page_break()

#     impact_table = doc.add_table(rows=4, cols=3)
#     impact_table.style = "Table Grid"
#     set_table_borders(impact_table)

#     impact_table.rows[0].cells[0].text = "RISK"
#     impact_table.rows[0].cells[1].text = "IMPACT"
#     impact_table.rows[0].cells[2].text = "ACTION"

#     impact_table.rows[1].cells[0].text = "MEDIUM"
#     impact_table.rows[2].cells[0].text = "HIGH"
#     impact_table.rows[3].cells[0].text = "VERY HIGH"

#     set_cell_color(impact_table.rows[1].cells[0], "FFFF00")
#     set_cell_color(impact_table.rows[2].cells[0], "FFA500")
#     set_cell_color(impact_table.rows[3].cells[0], "FF0000")

#     # 🔷 SIGNATURE
#     doc.add_paragraph("\nDirector I/C        DUTY OFFICER")

#     output_file = "FINAL_IMD_BULLETIN.docx"
#     doc.save(output_file)

#     return output_file